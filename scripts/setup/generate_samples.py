import os
from pathlib import Path
from docx import Document

def generate_samples():
    sample_dir = Path("data/sample")
    sample_dir.mkdir(parents=True, exist_ok=True)
    
    # 1. Generate sample_pmo_update.docx
    doc = Document()
    doc.add_heading("PMO Weekly Update: Project Apex ERP Transformation", level=1)
    
    doc.add_heading("1. Executive Summary", level=2)
    doc.add_paragraph(
        "Project Apex ERP Transformation is currently at yellow status due to delays in third-party API integrations. "
        "The core ledger customization is 90% complete, and data migration testing for Phase 1 has reached 85% success rate. "
        "However, external middleware delivery by partner VendorX is 3 weeks overdue, impacting the downstream testing schedule."
    )
    
    doc.add_heading("2. Key Milestones", level=2)
    table_milestones = doc.add_table(rows=1, cols=3)
    hdr_cells = table_milestones.rows[0].cells
    hdr_cells[0].text = 'Milestone'
    hdr_cells[1].text = 'Target Date'
    hdr_cells[2].text = 'Status'
    
    milestones = [
        ("Core Customization", "2026-05-15", "Completed"),
        ("Middleware Integration", "2026-06-01", "Delayed"),
        ("UAT Phase 1", "2026-06-15", "At Risk")
    ]
    for ms, dt, st in milestones:
        row_cells = table_milestones.add_row().cells
        row_cells[0].text = ms
        row_cells[1].text = dt
        row_cells[2].text = st
        
    doc.add_heading("3. Risks, Actions, Issues, and Dependencies (RAID)", level=2)
    doc.add_paragraph(
        "We are tracking several items that require immediate attention from the leadership team:"
    )
    
    # Add RAID detail paragraph for parser
    doc.add_paragraph(
        "RISK: Late API delivery from VendorX. Due to resource constraints at VendorX, the API endpoints for inventory sync "
        "are delayed. Impact: High. Mitigation: Daily standups with VendorX and preparation of stub services."
    )
    
    doc.add_paragraph(
        "ISSUE: Test environment database crash. The QA database instance crashed on Friday, causing 2 days of UAT downtime. "
        "Status: Resolved. Action taken: Restored from backup and increased DB storage allocation."
    )
    
    doc.add_paragraph(
        "DEPENDENCY: Database schema freeze by DB Admin team. We cannot finalize the API mappings until the DB schema is frozen "
        "by the infrastructure team. Due Date: 2026-05-30."
    )
    
    doc.add_heading("4. Escalations", level=2)
    doc.add_paragraph(
        "ESCALATION: VendorX SLA breach on Phase 2 deliverables. VendorX has missed their third consecutive delivery checkpoint. "
        "We request formal contract penalty clause execution and escalation of this issue. Routing Target: Steering Committee."
    )
    
    doc.save(sample_dir / "sample_pmo_update.docx")
    print("Generated sample_pmo_update.docx successfully.")
    
    # 2. Generate sample_governance_report.txt
    txt_content = """PROJECT GOVERNANCE REVIEW: PROJECT HELIOS CLOUD MIGRATION
Status: RED
Phase: Assessment & Planning
Budget Variance: $120,000 overrun due to unexpected storage egress costs and extended VM provisioning.
Assessment Date: 2026-05-20
Reviewer: Lead PMO Analyst

1. OVERVIEW
Project Helios is the migration of legacy financial core systems to the AWS cloud. The project is currently experiencing severe headwinds due to vendor staffing issues and unbudgeted cloud infrastructure expenses.

2. DETAILED STATUS & ISSUES
- BUDGET OVERRUN: The project has incurred an unexpected $120,000 cost overrun. This is driven by legacy systems transmitting redundant test data payloads, resulting in high network egress charges.
- VENDOR STAFFING DEVIATION: Vendor did not allocate certified cloud migration engineers as specified in the Master Services Agreement (MSA). This staffing gap has led to a 2-week slip in the migration architecture sign-off.
- DATA SECURITY GAP: The compliance team has raised a concern that PII data encryption keys are currently managed under a shared service account rather than dedicated IAM roles.

3. ESCALATION ROUTING
- ESCALATION: We request contract penalty clause enforcement due to the vendor's failure to supply certified engineers. This breach must be resolved immediately. Routing Target: PMO Lead.
- ESCALATION: Approval required for an additional budget allocation of $50,000 to cover transit gateways and network traffic logging. Routing Target: Steering Committee.

4. NEXT STEPS
- Implement traffic filtering to reduce network egress costs by May 28.
- Finalize IAM separation of duties by June 5.
"""
    
    with open(sample_dir / "sample_governance_report.txt", "w", encoding="utf-8") as f:
        f.write(txt_content)
    print("Generated sample_governance_report.txt successfully.")

if __name__ == "__main__":
    generate_samples()
