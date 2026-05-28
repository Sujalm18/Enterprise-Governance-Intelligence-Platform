import logging
import io
import re
import time
from pathlib import Path
from pypdf import PdfReader
from docx import Document as DocxDocument
from PIL import Image
import fitz  # PyMuPDF

logger = logging.getLogger("governance_copilot.ingestion.parser")


# Safeguards
MAX_OCR_PAGES = 50

# EasyOCR Lazy-Loaded Reader Singleton
_OCR_READER = None

def _get_ocr_reader():
    """Initializes EasyOCR Reader once and caches it in memory."""
    global _OCR_READER
    if _OCR_READER is None:
        logger.info("Initializing EasyOCR English Reader (loading detection & recognition models)...")
        import easyocr
        # Initialize easyocr reader with English, gpu disabled by default unless cuda available
        # easyocr automatically uses GPU if gpu=True (which is default). Let's explicitly specify gpu=False
        # to ensure 100% stability across CPU-only Windows configurations, or let easyocr handle it.
        # We'll set gpu=False to be extremely safe, consistent with prompt requests.
        _OCR_READER = easyocr.Reader(['en'], gpu=False)
    return _OCR_READER

def parse_file(file_path: str, file_type: str) -> str:
    """
    Parses PDF, DOCX, or TXT documents and extracts text content.
    No OCR or image recognition is performed unless scanned PDFs are detected.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found at {file_path}")
    
    file_type = file_type.lower().strip()
    logger.info(f"Parsing file: {path.name} of type: {file_type}")
    
    try:
        if file_type == "pdf" or path.suffix.lower() == ".pdf":
            parsed = parse_pdf(path)
        elif file_type in ("docx", "doc") or path.suffix.lower() in (".docx", ".doc"):
            parsed = parse_docx(path)
        elif file_type == "txt" or path.suffix.lower() == ".txt":
            parsed = parse_txt(path)
        elif file_type == "xlsx" or path.suffix.lower() == ".xlsx":
            parsed = parse_xlsx(path)
        else:
            # Fallback to text parsing
            logger.warning(f"Unknown type '{file_type}'. Attempting text parsing.")
            parsed = parse_txt(path)
        return f"SOURCE_FILE: {path.name}\n{parsed}"
    except Exception as e:
        logger.error(f"Failed to parse {path.name}: {e}")
        raise ValueError(f"Error parsing file {path.name}: {str(e)}")

def parse_pdf(path: Path) -> str:
    """
    Attempts searchable text extraction. If text is empty or too short (< 50 chars),
    triggers OCR fallback parsing using EasyOCR.
    """
    reader = PdfReader(path)
    text_content = []
    
    for idx, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text_content.append(page_text)
            
    extracted_text = "\n\n".join(text_content)
    
    quality = detect_poor_pdf_extraction(extracted_text, len(reader.pages))
    # Trigger OCR Fallback if text extracted is empty, too short, sparse, or visibly noisy.
    if quality["confidence"] < 0.45:
        logger.info(
            "PDF standard extraction quality is poor "
            f"(confidence={quality['confidence']:.2f}, noise={quality['noise_score']:.2f}, "
            f"density={quality['text_density']:.2f}, malformed={quality['malformed_ratio']:.2f}). "
            "Invoking EasyOCR fallback..."
        )
        try:
            return parse_pdf_ocr(path)
        except Exception as ocr_err:
            logger.error(f"EasyOCR fallback failed for {path.name}: {ocr_err}")
            if extracted_text.strip():
                logger.warning("Using imperfect pypdf extraction because OCR fallback failed.")
                return extracted_text
            raise ValueError(f"PDF file returned empty text. OCR fallback also failed: {str(ocr_err)}")
        
    return extracted_text


def detect_poor_pdf_extraction(text: str, page_count: int = 1) -> dict:
    """Scores pypdf extraction quality to decide whether OCR enhancement is needed."""
    stripped = text.strip()
    chars = len(stripped)
    if chars == 0:
        return {"confidence": 0.0, "noise_score": 1.0, "text_density": 0.0, "malformed_ratio": 1.0}

    lines = [line.strip() for line in stripped.splitlines() if line.strip()]
    alpha_num = sum(ch.isalnum() for ch in stripped)
    symbols = sum((not ch.isalnum() and not ch.isspace()) for ch in stripped)
    suspicious_tokens = len(re.findall(r"\b(?:govemance|esca1ation|dependencles|rn|lssue|r1sk)\b", stripped, re.IGNORECASE))
    broken_words = len(re.findall(r"\b[A-Za-z](?:\s+[A-Za-z]){3,}\b", stripped))
    malformed_ratio = min(1.0, (suspicious_tokens + broken_words) / max(1, len(stripped.split())))
    symbol_ratio = symbols / max(1, chars)
    line_coherence = sum(1 for line in lines if len(line) >= 20) / max(1, len(lines))
    text_density = chars / max(1, page_count * 1200)
    short_multi_page = page_count > 1 and chars < page_count * 250
    noise_score = min(1.0, symbol_ratio * 2.5 + malformed_ratio + (0.25 if line_coherence < 0.35 else 0))
    confidence = 1.0
    if chars < 50 or short_multi_page:
        confidence -= 0.6
    if text_density < 0.08:
        confidence -= 0.25
    if noise_score > 0.35:
        confidence -= 0.35
    confidence = max(0.0, min(1.0, confidence))
    return {
        "confidence": confidence,
        "noise_score": noise_score,
        "text_density": text_density,
        "malformed_ratio": malformed_ratio,
    }

def parse_pdf_ocr(path: Path) -> str:
    """
    Converts PDF pages into in-memory images using PyMuPDF and extracts text using EasyOCR.
    Includes safeguards for page count limits and logs rich execution performance metrics.
    """
    import numpy as np
    
    start_time = time.time()
    logger.info(f"OCR started for {path.name}")
    
    try:
        doc = fitz.open(path)
        page_count = len(doc)
        logger.info(f"OCR page count: {page_count}")
        
        # Guardrail page count limit check
        if page_count > MAX_OCR_PAGES:
            raise ValueError(f"Scanned PDF page count ({page_count}) exceeds safe guardrail limit of {MAX_OCR_PAGES} pages.")
            
        reader = _get_ocr_reader()
        extracted_text = []
        confidences = []
        
        for page_num in range(page_count):
            logger.info(f"Performing OCR on page {page_num + 1}/{page_count}")
            page = doc.load_page(page_num)
            pix = page.get_pixmap(alpha=False)
            
            # EasyOCR officially supports numpy.ndarray input. Convert the rendered
            # PDF page into a contiguous RGB ndarray before handing it to readtext.
            img_arr = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
            
            if pix.n == 1:
                img_arr = np.repeat(img_arr, 3, axis=2)
            elif pix.n > 3:
                img_arr = img_arr[:, :, :3]
            img_arr = np.ascontiguousarray(img_arr, dtype=np.uint8)
            logger.debug(f"EasyOCR input type for page {page_num + 1}: {type(img_arr).__name__}, shape={img_arr.shape}, dtype={img_arr.dtype}")
            
            # OCR using reader with detailed details enabled to retrieve confidence scores
            results = reader.readtext(img_arr, detail=1)
            
            page_text_blocks = []
            for bbox, text, confidence in results:
                page_text_blocks.append(text)
                confidences.append(confidence)
                
            page_text = " ".join(page_text_blocks)
            if page_text.strip():
                extracted_text.append(page_text)
                
        duration = time.time() - start_time
        full_text = "\n\n".join(extracted_text)
        char_count = len(full_text)
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
        
        logger.info(f"OCR completed for {path.name}")
        logger.info(f"OCR duration: {duration:.2f} seconds")
        logger.info(f"OCR page count: {page_count}")
        logger.info(f"OCR character count: {char_count}")
        logger.info(f"OCR confidence: {avg_confidence:.2%}")
        
        if not full_text.strip():
            raise ValueError("EasyOCR completed processing but failed to extract any readable alphanumeric content.")
            
        return full_text
        
    except Exception as e:
        logger.error(f"Error during OCR parsing of {path.name}: {e}")
        raise ValueError(f"Failed to parse scanned PDF with OCR: {str(e)}")


def parse_docx(path: Path) -> str:
    doc = DocxDocument(path)
    text_content = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text_content.append(paragraph.text)
            
    # Also parse tables if present (useful for status updates)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
            if row_text:
                text_content.append(row_text)
                
    if not text_content:
        raise ValueError("DOCX file returned empty text.")
        
    return "\n".join(text_content)

def parse_txt(path: Path) -> str:
    # Try common encodings
    encodings = ["utf-8", "latin-1", "cp1252"]
    for encoding in encodings:
        try:
            with open(path, "r", encoding=encoding) as f:
                content = f.read()
                return content
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Unable to decode text file {path.name} with standard encodings.")


def parse_xlsx(path: Path) -> str:
    """Extracts workbook text while preserving sheet, headers, and row structure."""
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=True)
    text_blocks = []
    for sheet in workbook.worksheets:
        text_blocks.append(f"SHEET: {sheet.title}")
        rows = []
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value).strip() for value in row]
            if any(values):
                rows.append(values)
        if not rows:
            continue
        headers = rows[0]
        text_blocks.append("HEADERS: " + " | ".join(headers))
        for values in rows[1:]:
            pairs = []
            for header, value in zip(headers, values):
                if header and value:
                    pairs.append(f"{header}: {value}")
            if pairs:
                text_blocks.append("ROW: " + " | ".join(pairs))
    full_text = "\n".join(text_blocks)
    if not full_text.strip():
        raise ValueError("XLSX file returned empty text.")
    return full_text
